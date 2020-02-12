

'''

{
  "DCRI": {
    "nome": "CRI",
    "obiettivi_formativi": "Il corso di formazione ",
    "competenze_in_uscita": "Si prevede che il partecipante",
    "lezioni": {
      "1": {
        "lezione": "Introduzione al Movimento Internazionale di Croce Rossa e Mezzaluna Rossa",
        "id": "1IAMIDCREMR",
        "argomento": "La nascita di unidea:",
        "ore": "2"
      },
    }
  }
}

'''


from docx import Document
import json
import os

startInputPath = './file_'
startOutputPath = './output'


def parseDocx(document, path_save):
    dict = {}
    isCaratteristiche = True
    isTitle = True
    indiceLezione = 1
    sigla = ''

    count = 0

    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            # Table caratteristiche
            if isCaratteristiche:
                if sigla == '':
                    sigla = cells[1].text.split('\n')[-1].strip()
                    dict[sigla] = {}
                    dict[sigla]['lezioni'] = {}
                if isTitle:
                    dict[sigla]['nome'] = ''.join(cells[1].text.split('\n')[:-1])
                    isTitle = False
                # Campi
                else:
                    if cells[0].text.strip().lower() == 'obiettivi formativi':
                        if len(cells) == 2:
                            dict[sigla]['obiettivi_formativi'] = cells[1].text.replace('\n', '').replace('§', '-')
                        else:
                            dict[sigla]['obiettivi_formativi'] = cells[2].text.replace('\n', '').replace('§', '-')
                    elif cells[0].text.strip().lower() == 'output competenze' or \
                            cells[0].text.strip().lower() == 'output\ncompetenze' or \
                            cells[0].text.strip().lower() == 'output \ncompetenze':
                        if len(cells) == 2:
                            dict[sigla]['competenze_in_uscita'] = cells[1].text.replace('\n', '').replace('§', '-')
                        else:
                            dict[sigla]['competenze_in_uscita'] = cells[2].text.replace('\n', '').replace('§', '-')
                    elif cells[0].text.strip().lower().replace('\n', '') == 'verifica e valutazione':
                        print(cells[1].text)

            # Table lezioni
            else:
                # Riga carateristiche lezione
                if len(cells) == 6 and cells[1].text != '':
                    if not cells[1].text.strip().lower() == 'lezioni':
                        if cells[0].text.replace('\n', ''):

                            dict[sigla]['lezioni'].update({
                                indiceLezione: {
                                    "lezione": cells[1].text.replace('\n', '').replace('§', '-'),
                                    "id": "{}{}".format(
                                        str(indiceLezione),
                                        ''.join(
                                            [i[0] for i in cells[1].text.replace('\n', '').split()]
                                        )
                                    ).upper(),
                                    "argomento": cells[2].text.replace('\n', '').replace('§', '-'),
                                    "ore": cells[4].text.replace('\n', '') if not 'TOT' in cells[4].text.replace('\n', '') else cells[4].text.strip().replace('\n', '').replace('\t', '').split('TOT')[1].strip()
                                },
                            })
                            indiceLezione += 1
                else:
                    print('else')
                    # # PARSE DOCX Formattati male
                    # if count < 2:
                    #     count += 1
                    #     continue
                    # print(len(cells))
                    # print('lezione', cells[3].text.replace('\n', '').replace('§', '-'))
                    # print('argomento', cells[5].text.replace('\n', '').replace('§', '-'))
                    # print('ore', cells[6].text.replace('\n', '').replace('§', '-'))
                    # print('\n\n')
                    # if len(cells) == 8:
                    #
                    #     dict[sigla]['lezioni'].update({
                    #         indiceLezione: {
                    #             "lezione": cells[1].text.replace('\n', '').replace('§', '-'),
                    #             "id": "{}{}".format(
                    #                 str(indiceLezione),
                    #                 ''.join(
                    #                     [i[0] for i in cells[1].text.replace('\n', '').split()]
                    #                 )
                    #             ).upper(),
                    #             "argomento": cells[5].text.replace('\n', '').replace('§', '-'),
                    #             "nota": "",
                    #             "ore": cells[6].text.replace('\n', '')
                    #         },
                    #     })
                    #     indiceLezione += 1

        isCaratteristiche = False

    os.makedirs(os.path.dirname(path_save), exist_ok=True)
    with open(path_save, 'w') as f:
        json.dump(dict, f)


def parseDocxToCSV(document, path_save):
    list = []

    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells):
                if cells[0].text.strip().lower().replace(' ', '').replace('\n', '') == 'verificaevalutazione':
                    for el in cells[1:]:
                        if el.text:
                            list.append(
                                {
                                    "nome": path_save.split('/')[-1].split('.')[0],
                                    "value": el.text
                                }
                            )
            else:
                list.append(
                    {
                        "nome": path_save.split('/')[-1].split('.')[0],
                        "value": ""
                    }
                )


    return list


def parsRec(path_start='./file', path_save='./output'):
    list = []
    if os.path.isfile(path_start) and path_start.split('.')[-1] == 'docx':
        print('**[{}]'.format(path_start))
        document = Document(path_start)
        # parseDocx(document, path_save)
        list.extend(parseDocxToCSV(document, path_save))
    elif os.path.isdir(path_start):
        for el in os.listdir(path_start):
            list.extend(
                parsRec(
                    path_start=os.path.join(path_start, el),
                    path_save=os.path.join(path_save, el).replace('docx', 'json').replace('doc', 'json')
                )
            )
    else:
        print('{}'.format(path_start))

    return list


if __name__ == "__main__":
    l = parsRec()

    os.makedirs(os.path.dirname('./output/file.json'), exist_ok=True)
    with open('./output/file.json', 'w') as f:
        json.dump(l, f)


